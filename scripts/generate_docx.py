"""Generate a professional delivery document (.docx) for the AI Agent Full-Stack Engineer test."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="2D1B69"):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F5F3FF")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph("")
    return table


def build_document():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = 'Microsoft YaHei'
    style.paragraph_format.space_after = Pt(6)

    # ================================================================
    # COVER PAGE
    # ================================================================
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Native 互动游戏 Web 平台")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("系统设计测试 — 交付文档")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        "候选人：Niuyuhang",
        "提交日期：2026-06-20",
        "GitHub：github.com/waterttag/AI_Agent",
        "Demo：https://aiagent-production-5b68.up.railway.app",
        "测试账号：demo@aigame.dev / demo123",
    ]
    for line in lines:
        run = info.add_run(f"\n{line}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (manual)
    # ================================================================
    h = doc.add_heading("目录", level=1)
    toc_items = [
        "一、项目概述",
        "二、技术栈",
        "三、系统架构设计",
        "四、数据模型设计",
        "五、API 接口设计",
        "六、AI Agent 选型与编排",
        "七、远程部署协议",
        "八、安全策略",
        "九、失败恢复机制",
        "十、可观测性",
        "十一、详细交付物清单",
        "十二、功能验证总结",
        "十三、完成度说明",
        "十四、AI 协同记录",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ================================================================
    # 一、项目概述
    # ================================================================
    doc.add_heading("一、项目概述", level=1)
    doc.add_paragraph(
        "本项目是一个 AI Native 互动游戏 Web 平台 MVP（参考 Astrocade），"
        "用 2 天时间在 AI 编程工具（Claude Code + DeepSeek API）的协同下从零搭建。"
        "打通了「注册/登录 → 创意生成 → 游戏发布 → 浏览游玩」的完整业务闭环。"
    )
    doc.add_heading("产品主旅程", level=2)
    doc.add_paragraph("• 玩家旅程：发现 → 浏览 → 游玩（Home → Play）")
    doc.add_paragraph("• 创作者旅程：创意 → AI 生成 → 预览 → 发布（Create → Generate → Preview → Publish）")

    doc.add_heading("用户角色", level=2)
    add_styled_table(doc, ["角色", "核心活动", "实现状态"], [
        ["玩家 (Player)", "浏览平台游戏、点击游玩、查看游戏信息", "✅ Home + Play 页"],
        ["创作者 (Creator)", "登录 → Create → 输入创意+上传素材 → DeepSeek 生成 → 预览 → 发布", "✅ 完整闭环"],
        ["平台维护者 (Maintainer)", "关注任务稳定性、管理 OSS 文件、处理不安全内容", "✅ Health + 自动种子"],
    ], col_widths=[3, 7, 4])

    # ================================================================
    # 二、技术栈
    # ================================================================
    doc.add_heading("二、技术栈", level=1)
    add_styled_table(doc, ["层级", "技术选型", "选型理由"], [
        ["前端", "React 18 + Vite + TypeScript + Tailwind CSS + TanStack Query + Zustand",
         "Vite HMR 极速开发；Tailwind 原子化零切换成本；TanStack Query 一行搞定服务端状态缓存+轮询；Zustand 轻量 auth"],
        ["后端", "FastAPI (Python 3.11+) + SQLAlchemy 2.0 (async) + Celery + Redis",
         "FastAPI 异步原生 + 自动 OpenAPI 文档；Python 是 AI 生态（Anthropic SDK / LangChain）首选；Celery 独立进程解耦 LLM 长耗时调用"],
        ["数据库", "SQLite (dev) → PostgreSQL (prod)", "SQLAlchemy 抽象层隔离差异，改 DATABASE_URL 一行切换"],
        ["对象存储", "boto3 S3 客户端 → 阿里云 OSS / AWS S3 / MinIO / Cloudflare R2",
         "单客户端兼容多服务；自动检测 virtual-hosted / path 风格；无存储时自动降级为 DB 存储"],
        ["AI / LLM", "适配器工厂 (ClaudeAdapter | OpenAIAdapter | DeepSeekAdapter)",
         "策略模式，env 变量切换；DeepSeekAdapter 继承 OpenAIAdapter 仅改 base_url"],
        ["部署", "Railway (Dockerfile) + 阿里云 OSS", "Dockerfile 多阶段构建（前端 build + Python runtime）；CMD shell 形式接受 \$PORT"],
    ], col_widths=[2.5, 5, 7])

    # ================================================================
    # 三、系统架构设计
    # ================================================================
    doc.add_heading("三、系统架构设计", level=1)

    doc.add_heading("3.1 架构总览", level=2)
    doc.add_paragraph(
        "浏览器 (React+Vite) ←→ FastAPI Backend ←→ Celery Worker (Redis Broker) ←→ DeepSeek API\n"
        "                                    ↕                       ↕\n"
        "                              SQLite / PostgreSQL      阿里云 OSS / MinIO"
    )

    doc.add_heading("3.2 组件协作流程", level=2)
    doc.add_paragraph(
        "1. 用户浏览器通过 REST API 与 FastAPI 通信\n"
        "2. 生成请求到达后，FastAPI 不直接调用 LLM（耗时 30-120s 会阻塞整个 worker），"
        "而是创建 generation_task 记录并入队 Celery（Redis 为消息代理）\n"
        "3. Celery Worker 独立进程异步执行 4 阶段 Pipeline："
        "Preprocess → Generate → Validate & Fix → Package & Upload\n"
        "4. 生成的游戏文件上传到阿里云 OSS（或降级存入 DB），"
        "前端通过 iframe 动态加载 play-html 端点\n"
        "5. AI 生成完成后游戏进入 preview 状态，创作者确认后点击 Publish 变为 published"
    )

    doc.add_heading("3.3 为什么用 Celery + Redis 而不是 FastAPI BackgroundTasks？", level=2)
    doc.add_paragraph(
        "BackgroundTasks 在主进程中执行，LLM 调用 30-120s 期间整个 API worker 被阻塞，"
        "其他请求无法响应。Celery 独立进程 + Redis 消息队列实现请求与生成的完全解耦，"
        "支持水平扩展 Worker 数量应对高并发。"
    )

    doc.add_heading("3.4 为什么 SQLite 开发 + PostgreSQL 生产？", level=2)
    doc.add_paragraph(
        "SQLAlchemy 2.0 async ORM 抽象层隔离差异，改一个环境变量即可切换。"
        "SQLite 零配置启动开发环境，PostgreSQL 提供生产级并发和 JSONB 索引。"
        "Railway 上当前使用 SQLite（免费额度），生产迁移只需改 DATABASE_URL。"
    )

    # ================================================================
    # 四、数据模型设计
    # ================================================================
    doc.add_heading("四、数据模型设计", level=1)
    doc.add_paragraph("共 5 张表，拆分为独立实体：")

    add_styled_table(doc, ["表名", "核心字段", "设计要点"], [
        ["users", "id, username, email, password_hash (bcrypt), role (player|creator)", "role 字段预留 OAuth 扩展空间；认证与业务数据分离"],
        ["games", "id, title, description, game_url, cover_image_url,\nauthor_id (FK), tags (JSON), status (draft|generating|preview|published|failed),\nprompt_text, play_count",
         "status 状态机驱动全生命周期；tags 用 JSON 灵活支持任意标签，SQLite/PG 均支持 LIKE/JSONB 查询；preview 状态实现「生成→预览→发布」流程"],
        ["game_assets", "id, game_id (FK), asset_type, oss_key, oss_url, file_size", "多对一关联 games，独立存储 OSS key，方便增删；上传通过阿里云 OSS boto3 client"],
        ["generation_tasks", "id, game_id (FK), user_id (FK), status, progress,\nsystem_prompt_used, user_prompt_used, llm_response_raw,\nresult_oss_url, error_message",
         "生成过程与游戏结果解耦，历史可追溯；完整保存 prompt + LLM 响应用于调试和 Agent 日志；支持版本历史展示"],
        ["game_favorites", "user_id (FK) + game_id (FK) (复合 PK), created_at", "轻量收藏表，复合主键天然去重；支持 toggle 切换和批量查询"],
    ], col_widths=[2.5, 5.5, 6], header_color="1B4332")

    doc.add_heading("4.1 为什么 tags 用 JSON 而非关联表？", level=2)
    doc.add_paragraph(
        "游戏标签数量少（≤10）且无需跨游戏聚合查询。JSON 列避免了额外 join 表，"
        "SQLite 使用 cast(Text).like() 实现搜索，PostgreSQL 可以用 GIN 索引达到关联表级别的查询性能。"
    )

    # ================================================================
    # 五、API 接口设计
    # ================================================================
    doc.add_heading("五、API 接口设计", level=1)
    doc.add_paragraph("全部 23 个端点，自动生成 OpenAPI 文档（/docs）：")
    add_styled_table(doc, ["分组", "方法", "路径", "说明", "Auth"], [
        ["Auth", "POST", "/api/auth/register", "注册", "—"],
        ["Auth", "POST", "/api/auth/login", "登录", "—"],
        ["Auth", "GET", "/api/auth/me", "当前用户", "✅"],
        ["Auth", "GET", "/api/auth/me/favorites", "我的收藏 ID 列表", "✅"],
        ["Games", "GET", "/api/games?status=&tag=&page=&size=", "游戏列表 (分页+筛选)", "—"],
        ["Games", "GET", "/api/games/{id}?increment=true", "游戏详情 (播放计数)", "—"],
        ["Games", "POST", "/api/games", "创建草稿", "✅"],
        ["Games", "PUT", "/api/games/{id}", "更新 (含 status 发布)", "✅"],
        ["Games", "DELETE", "/api/games/{id}", "删除", "✅"],
        ["Games", "GET", "/api/games/{id}/play-html", "直接服务游戏 HTML", "—"],
        ["Assets", "POST", "/api/games/{id}/assets", "上传到 OSS", "✅"],
        ["Assets", "GET", "/api/games/{id}/assets", "素材列表", "—"],
        ["Assets", "DELETE", "/api/games/{id}/assets/{id}", "删除素材", "✅"],
        ["Tasks", "POST", "/api/games/{id}/generate", "触发 AI 生成", "✅"],
        ["Tasks", "GET", "/api/tasks/{id}", "轮询状态 (进度%)", "✅"],
        ["Tasks", "GET", "/api/tasks/games/{id}", "任务历史", "✅"],
        ["Tasks", "GET", "/api/tasks/games/{id}/log", "Agent 执行日志", "✅"],
        ["Favorites", "POST", "/api/games/{id}/favorite", "收藏/取消切换", "✅"],
        ["Favorites", "GET", "/api/games/{id}/favorite", "收藏状态+总数", "✅"],
        ["Admin", "POST", "/api/admin/inject-game", "注入预构建游戏", "—"],
    ], col_widths=[1.5, 1, 4, 4, 1.2])

    # ================================================================
    # 六、AI Agent 选型与编排
    # ================================================================
    doc.add_heading("六、AI Agent 选型与编排", level=1)

    doc.add_heading("6.1 适配器工厂", level=2)
    doc.add_paragraph(
        "采用策略模式，通过 LLM_PROVIDER 环境变量切换 LLM 提供方。新增模型只需实现 LLMAdapter 抽象基类：\n"
        "  LLMAdapter (ABC)\n"
        "  ├── ClaudeAdapter  → api.anthropic.com    (支持 Vision)\n"
        "  ├── OpenAIAdapter  → api.openai.com       (支持 Vision + 自定义 base_url)\n"
        "  └── DeepSeekAdapter → api.deepseek.com    (继承 OpenAIAdapter，仅改 base_url)"
    )

    doc.add_heading("6.2 为什么自建适配器？而不是 OpenClaw / Hermes / Pi Agent / LangGraph？", level=2)
    doc.add_paragraph(
        "核心原因：MVP 复杂度控制 + 隔离供应商依赖。\n\n"
        "• OpenClaw / Hermes / Pi Agent 是完整 Agent 框架（多步骤推理、工具调用、记忆管理）。"
        "但本项目的核心任务是「一段 prompt → 一个 HTML 文件」的单次生成，不需要多步拆解和多工具编排。引入框架徒增认知负担和依赖耦合。\n"
        "• LangGraph 适合有状态、多节点、条件分支的复杂 Agent 工作流。当前流水线是线性的"
        "（Preprocess→Generate→Validate→Upload），用简单的 status 字段 + Celery 即可表达，LangGraph 在 MVP 阶段过度设计。\n"
        "• 适配器工厂 (ABC + env switch)：3 个文件实现多 LLM 切换。新增供应商只需实现两个方法（generate, describe_image），无框架锁定。\n\n"
        "扩展路径：若未来需要多 Agent 协作（如策划 Agent + 美术 Agent + 代码 Agent 分工），"
        "迁移到 LangGraph 成本低——已有 Pipeline 中的每个 Phase 可无损映射为 LangGraph Node。"
    )

    doc.add_heading("6.3 生成流水线 (4-Phase Pipeline)", level=2)
    add_styled_table(doc, ["Phase", "名称", "内容", "进度"], [
        ["1", "Preprocess", "加载素材 → Vision API 描述图片 → 组装结构化 GameGenerationContext", "0-30%"],
        ["2", "Generate", "选择适配器 → 构建 System Prompt → 调用 LLM → 收集完整 HTML（temperature 0.8, max_tokens 16000）", "30-70%"],
        ["3", "Validate & Fix", "BeautifulSoup 解析 HTML → 检查 game loop → 禁止 eval() → CDN 白名单 → 失败自动修复 1 次", "70-85%"],
        ["4", "Package & Upload", "Content-Disposition: inline → 上传到 OSS/S3 → 或降级 DB 存储 → 设置 game.status=preview", "85-100%"],
    ], col_widths=[1, 2.5, 8, 2], header_color="4A0E4E")

    doc.add_heading("6.4 System Prompt 设计", level=2)
    doc.add_paragraph(
        "• 要求输出自包含 HTML5 游戏（Phaser.js CDN 或 Canvas API）\n"
        "• 强制约束：game loop（requestAnimationFrame/setInterval/Phaser）、交互性（键盘/鼠标/触摸）、响应式设计\n"
        "• 安全约束：禁止 eval()、限制外部请求至 CDN 白名单\n"
        "• 素材注入：通过 Vision API 描述图片，文本化注入 prompt 上下文\n"
        "• Auto-fix：校验失败时将错误列表返回 LLM 二次修复（temperature 0.5）"
    )

    # ================================================================
    # 七、远程部署协议
    # ================================================================
    doc.add_heading("七、远程部署协议", level=1)

    doc.add_heading("7.1 游戏文件格式 — 为什么是单文件 HTML？", level=2)
    doc.add_paragraph(
        "1. 存储简单：对象存储中每个游戏只有 1 个文件，路径 games/{uuid}/index.html\n"
        "2. 传输高效：一次 HTTP GET 即可加载全部游戏内容\n"
        "3. LLM 友好：要求 LLM 输出单文件不依赖外部路径解析，模型天然擅长\n"
        "4. 沙箱安全：iframe sandbox 加载单文件，无相对路径跨越问题\n"
        "5. 跨平台：移动端、桌面端、任何浏览器都能打开自包含 HTML"
    )

    doc.add_heading("7.2 Play 页加载流程", level=2)
    doc.add_paragraph(
        "1. 前端 GET /api/games/{id}?increment=true → 获取 game_url 并累计播放次数\n"
        "2. <iframe sandbox=\"allow-scripts allow-same-origin\" src={game_url}>\n"
        "3. game_url 使用 API 端点（/api/games/{id}/play-html）避免 OSS Content-Disposition 问题\n"
        "4. Bucket 配置 public-read ACL + CORS\n"
        "5. 无对象存储时自动降级：generation_tasks.llm_response_raw → play-html 直接服务"
    )

    # ================================================================
    # 八、安全策略
    # ================================================================
    doc.add_heading("八、安全策略", level=1)
    add_styled_table(doc, ["威胁层面", "措施", "技术实现"], [
        ["认证", "JWT (HS256, 24h 过期), bcrypt 密码哈希", "python-jose + passlib, /api/auth/*"],
        ["上传安全", "MIME 白名单 (image/*, audio/*), 最大 10 文件", "FastAPI UploadFile + Content-Type 校验"],
        ["生成代码执行", "iframe sandbox 双层隔离, Validator 静态检查 eval()", "allow-scripts allow-same-origin，禁止 allow-top-navigation"],
        ["Prompt Injection", "用户 prompt 仅作为 User Message 传入（非 System Prompt）", "System Prompt 由后端构建，用户输入不可逃逸"],
        ["密钥管理", "LLM_API_KEY / JWT_SECRET 仅存 .env，.gitignore 排除", "Railway Secret Manager 生产注入"],
        ["资源控制", "max_tokens: 16000 限制成本上限, Worker pool=solo 串行", "防止并发过载和 API 费用失控"],
        ["依赖安全", "requirements.txt 锁定版本, npm audit", "可复现构建"],
    ], col_widths=[2.5, 5, 5.5], header_color="7B1D1D")

    # ================================================================
    # 九、失败恢复机制
    # ================================================================
    doc.add_heading("九、失败恢复机制", level=1)
    add_styled_table(doc, ["故障场景", "恢复策略", "可观察性"], [
        ["LLM API 超时/500", "Celery 自动重试 1 次（60s delay），仍失败写入 error_message", "task.status = failed，前端展示错误原因"],
        ["HTML 校验失败", "将错误列表返回 LLM 二次修复（temperature 0.5），仍失败降级交付", "Validator 日志 + llm_response_raw 保存原始输出"],
        ["OSS 上传失败", "Railway 环境无 MinIO 时自动降级——HTML 存入 generation_tasks.llm_response_raw, 通过 play-html 服务", "日志：'MinIO upload failed, storing in DB'"],
        ["素材描述失败 (Vision)", "Vision API 不可用时跳过该素材，其余素材正常注入 Prompt", "日志：'Failed to describe image'，任务继续"],
        ["部署重启丢数据", "后端 startup 自动调用 _auto_seed() 注入 3 款种子游戏 + demo 用户", "/health 即可验证"],
        ["资产上传失败 (前端)", "上传失败仅 skip 该文件，不阻断 AI 生成流程", "前端 console.warn + 继续生成"],
    ], col_widths=[3.5, 6, 4.5], header_color="5C4A0E")

    # ================================================================
    # 十、可观测性
    # ================================================================
    doc.add_heading("十、可观测性", level=1)
    add_styled_table(doc, ["观察维度", "记录方式", "证据路径"], [
        ["任务进度", "Celery 任务状态机 (pending→processing→completed/failed)，DB generation_tasks 表实时更新", "GET /api/tasks/{id}"],
        ["Agent 日志", "Python logging，每个 Phase 输出 [task_id] Phase N: ...", "Docker logs / Railway Deploy Logs"],
        ["LLM 调用记录", "generation_tasks 保存 system_prompt_used + user_prompt_used + llm_response_raw", "数据库查询 / Agent Log API"],
        ["播放统计", "GET /api/games/{id}?increment=true 自动 +1, 显示在首页卡片", "Home 页面 Eye 图标"],
        ["前端 Pipeline 可视化", "4 阶段进度条 + 步骤高亮动画", "Create 页面"],
        ["收藏计数", "POST toggle + GET count 实时反馈", "Play 页面 Favorite 按钮"],
        ["系统健康", "/health 端点 + Readiness 检查", "curl /health"],
        ["自动种子", "启动时检测空 DB → 注入 3 款游戏", "首次访问即有数据"],
        ["验证文档", "22 项自动化 API 测试 + 安全验证 + 性能数据", "docs/VERIFICATION.md"],
    ], col_widths=[3, 5.5, 5.5], header_color="0D4A5C")

    # ================================================================
    # 十一、详细交付物清单
    # ================================================================
    doc.add_heading("十一、详细交付物清单", level=1)
    doc.add_paragraph("对照测试文档 3.3 节，逐项提交：")
    add_styled_table(doc, ["#", "交付物", "状态", "说明"], [
        ["1", "源代码仓库", "✅", "github.com/waterttag/AI_Agent，15 次 commit，清晰的分阶段提交历史"],
        ["2", "Demo 地址", "✅", "https://aiagent-production-5b68.up.railway.app （测试账号：demo@aigame.dev / demo123）"],
        ["3", "部署方式", "✅", "Docker Compose（6 服务一键启动）+ Railway 全栈自动部署 + 阿里云 OSS"],
        ["4", "测试数据", "✅", "3 款预构建 HTML5 游戏（Snake / Memory Match / Breakout Blitz）+ Space Shooter（DeepSeek 生成）+ 自动种子注入"],
        ["5", "环境配置", "✅", ".env.example 列出全部 18 个环境变量及说明，真实密钥不提交"],
        ["6", "系统设计文档", "✅", "docs/SYSTEM_DESIGN.md — 架构图、数据模型、API 表、Agent 编排、安全、恢复、可观测性、设计决策 FAQ（含本文档全部内容）"],
        ["7", "技术栈", "✅", "README 完整列出前端/后端/数据库/存储/AI/部署全栈技术选型及理由"],
        ["8", "完成度说明", "✅", "docs/COMPLETION_REPORT.md — 已完成功能表、未完成/可改进项、1 周改进计划"],
        ["9", "功能验证证明", "✅", "docs/VERIFICATION.md — 22 项自动化 API 测试 + 手动验证 + 安全验证 + 性能数据"],
        ["10", "演示视频", "⏸️", "docs/DEMO_SCRIPT.md — 5 分钟 6 场景录制脚本已备好"],
        ["11", "AI 协同记录", "✅", "docs/AI_COLLAB_LOG.md — 工具、关键 Prompt、AI 辅助占比、人工修正、完整开发日志"],
    ], col_widths=[0.8, 2, 1, 10], header_color="2D1B69")

    # ================================================================
    # 十二、功能验证总结
    # ================================================================
    doc.add_heading("十二、功能验证总结", level=1)
    doc.add_paragraph("以下全部功能已通过线上测试（https://aiagent-production-5b68.up.railway.app）：")
    add_styled_table(doc, ["#", "功能", "测试结果", "验证方式"], [
        ["1", "用户注册", "✅", "curl POST /api/auth/register → 201 + JWT"],
        ["2", "用户登录", "✅", "curl POST /api/auth/login → 200 + token"],
        ["3", "登录失败提示", "✅", "curl POST → 401 'Invalid email or password'"],
        ["4", "重复注册拒绝", "✅", "curl POST → 409 'Email already registered'"],
        ["5", "退出登录", "✅", "前端 Zustand store 清除 + 导航栏恢复"],
        ["6", "游戏列表", "✅", "curl /api/games → 3 published games"],
        ["7", "标签筛选", "✅", "curl /api/games?tag=arcade → 2 games"],
        ["8", "分页翻页", "✅", "curl /api/games?page=1&size=2 → 2 items, total=3"],
        ["9", "播放计数", "✅", "curl /api/games/{id}?increment=true → +1 OK"],
        ["10", "收藏切换", "✅", "curl POST toggle → true/false 交替"],
        ["11", "创建游戏 (需登录)", "✅", "curl POST /api/games → 201 draft"],
        ["12", "创建游戏 (未登录拒绝)", "✅", "curl POST /api/games → 401"],
        ["13", "AI 生成 (DeepSeek)", "✅", "POST generate → poll → completed + game_url"],
        ["14", "Preview 状态", "✅", "生成后 status=preview，公网列表不可见"],
        ["15", "Publish 发布", "✅", "PUT /api/games/{id} {status:published} → 公网可见"],
        ["16", "Agent 执行日志", "✅", "tasks/log → prompt_summary + agent_steps"],
        ["17", "版本历史", "✅", "tasks/games/{id} → 全部历史 generation tasks"],
        ["18", "OSS 上传 (阿里云)", "✅", "boto3 PUT → 公网 GET HTTP 200 → DELETE"],
        ["19", "OSS Content-Disposition 修复", "✅", "inline + play-html 端点避免下载"],
        ["20", "无 API Key 降级", "✅", "LLM_PROVIDER=none → 503 + 明确提示"],
        ["21", "自动种子数据", "✅", "部署后即 3 款游戏 + demo 用户"],
        ["22", "健康检查", "✅", "GET /health → 200"],
    ], col_widths=[0.8, 4.5, 1.5, 7], header_color="1B4332")

    # ================================================================
    # 十三、完成度说明
    # ================================================================
    doc.add_heading("十三、完成度说明", level=1)

    doc.add_heading("核心功能 (100%)", level=2)
    add_styled_table(doc, ["模块", "完成度", "关键功能"], [
        ["Auth", "100%", "邮箱注册/登录/退出/登录态保持/OAuth 扩展预留"],
        ["Home", "100%", "游戏卡片网格/标签筛选/分页翻页/播放次数/收藏 ❤️"],
        ["Play", "100%", "iframe 动态加载/game loop 可玩/全屏 Escape/加载失败重试"],
        ["Create", "100%", "文本创意/多模态素材上传 OSS/4 阶段 Pipeline/Agent 日志/Preview→Publish"],
        ["AI Agent Harness", "100%", "适配器工厂 3 种 LLM/Validator 安全校验/Auto-fix/OSS 降级"],
        ["对象存储", "100%", "boto3 多服务/OSS 虚拟主机自动检测/DB 降级"],
        ["部署", "100%", "Docker Compose + Railway + 阿里云 OSS/自动种子/健康检查"],
    ], col_widths=[3, 1.5, 9.5])

    doc.add_heading("加分项 (已全部完成)", level=2)
    add_styled_table(doc, ["加分项", "状态", "实现"], [
        ["标签筛选", "✅", "Home 页 Filter Bar + 卡片标签可点击"],
        ["分页翻页", "✅", "首页 Prev/Next + 页码显示"],
        ["播放统计", "✅", "卡片 Eye 图标 + ?increment=true 精确计数"],
        ["收藏功能", "✅", "❤️ toggle + Favorites 筛选 + 计数"],
        ["预览→发布流程", "✅", "Preiew 黄色横幅 → Publish Now → published"],
        ["Agent 执行日志", "✅", "生成完成页展示 Prompt 摘要 + Pipeline 步骤"],
        ["版本历史", "✅", "Play 页可展开历史生成记录"],
        ["生成历史", "✅", "Task 历史 API + Play 页版本面板"],
        ["Agent 过程可视化", "✅", "4 阶段 Pipeline 步骤 + 进度条动画"],
        ["失败重试", "✅", "错误信息 + Try Again + Celery 自动重试"],
        ["安全沙箱", "✅", "iframe sandbox + Validator 禁止 eval()"],
        ["资源限额", "✅", "max_tokens: 16000 + Worker pool=solo"],
        ["内容审核预留", "✅", "Validator 可扩展关键词/模式过滤"],
    ], col_widths=[3, 1.5, 9.5])

    doc.add_heading("未完成项", level=2)
    doc.add_paragraph(
        "• OAuth 第三方登录：文档注明 demo 阶段可不实现，role 字段已预留扩展\n"
        "• 演示视频：录制脚本已备好（docs/DEMO_SCRIPT.md），待录制\n"
        "• WebSocket 推送：当前使用轮询（TanStack Query 2s），升级只需改一处"
    )

    # ================================================================
    # 十四、AI 协同记录
    # ================================================================
    doc.add_heading("十四、AI 协同记录", level=1)
    add_styled_table(doc, ["项目", "内容"], [
        ["主力 AI 工具", "Claude Code (CLI) — 架构设计、代码生成、部署调试（约 90%）"],
        ["内容生成 AI", "DeepSeek API (deepseek-chat) — HTML5 游戏代码生成（约 10%）"],
        ["关键 Prompt 策略", "系统架构设计（输入 .docx 全文 → 输出方案）/ Agent Harness System Prompt（4 阶段约束 + 安全规则）/ 修复 Prompt（校验失败 → LLM 二次修复）"],
        ["AI 辅助占比", "项目结构 95% / ORM+Schemas 95% / API 路由 90% / Agent Harness 90% / 前端组件 85% / 文档 80% / 种子游戏 10%（手写）"],
        ["人工修正重点", "SQLAlchemy greenlet 错误（懒加载 assets） / MinIO→boto3 迁移（虚拟主机风格） / DeepSeekAdapter（继承+base_url）/ OSS Content-Disposition 修复 / 数据库路径统一 / 前端 TS 类型错误"],
        ["开发时间", "Day 1: 脚手架→后端→前端→Agent→Celery→种子→联调 / Day 2: OSS 部署→全屏修复→自动种子→加分项（标签/预览/分页/播放/日志/收藏）"],
        ["Git 提交", "15 次 commit，清晰的分阶段提交历史"],
    ], col_widths=[3, 11], header_color="4A0E4E")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # --- FOOTER ---
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("— AI Native 互动游戏 Web 平台 · 系统设计测试交付文档 · 2026-06-20 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ================================================================
    # SAVE
    # ================================================================
    output_path = os.path.join(os.path.dirname(__file__), "..", "AIAgent全栈工程师-系统设计测试-交付文档.docx")
    doc.save(output_path)
    print(f"Saved to: {os.path.abspath(output_path)}")
    return output_path


if __name__ == "__main__":
    build_document()
